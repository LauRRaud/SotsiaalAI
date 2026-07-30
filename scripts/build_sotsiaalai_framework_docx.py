from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
LIGHT_GREY = "F2F4F7"
MID_GREY = "D0D5DD"
DARK_GREY = "475467"
RED = "B42318"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_borders(cell, color=MID_GREY, size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_repeat_header_footer(document: Document) -> None:
    document.settings.odd_and_even_pages_header_footer = True

    for section in document.sections:
        section.different_first_page_header_footer = True
        for header in (section.header, section.even_page_header):
            paragraph = header.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run("SotsiaalAI  |  organisatsioonikasutuse raamleping")
            run.font.name = "Calibri"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(DARK_GREY)

        for footer in (section.footer, section.even_page_footer):
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("Versioon 2026-07-30 · MUSTAND JURISTI ÜLEVAATUSEKS  |  ")
            run.font.name = "Calibri"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(DARK_GREY)
            add_page_field(paragraph)

        first_footer = section.first_page_footer
        first_paragraph = first_footer.paragraphs[0]
        first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = first_paragraph.add_run("SotsiaalAI OÜ · 2026-07-30 · kinnitamata tervikmustand")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(DARK_GREY)


def set_bottom_border(paragraph, color=BLUE, size="12", space="4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.33)
    section.footer_distance = Inches(0.34)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    styles["List Number"].font.name = "Calibri"
    styles["List Number"].font.size = Pt(11)
    styles["List Number"].paragraph_format.left_indent = Inches(0.5)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Number"].paragraph_format.space_after = Pt(8)
    styles["List Number"].paragraph_format.line_spacing = 1.167


def add_status_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(RED)
    run.font.all_caps = True


def add_metadata_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(DARK_GREY)


def add_clause_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.35)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.paragraph_format.widow_control = True
    match = re.match(r"^(\d+(?:\.\d+)+\.)\s+(.*)$", text)
    if match:
        number_run = paragraph.add_run(match.group(1) + " ")
        number_run.font.bold = True
        paragraph.add_run(match.group(2))
    else:
        paragraph.add_run(text)
    style_placeholders(paragraph)


def add_alpha_paragraph(document: Document, text: str, keep_with_next: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.55)
    paragraph.paragraph_format.first_line_indent = Inches(-0.3)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    match = re.match(r"^([a-zõäöü]\))\s+(.*)$", text, re.IGNORECASE)
    if match:
        marker = paragraph.add_run(match.group(1) + " ")
        marker.font.bold = True
        paragraph.add_run(match.group(2))
    else:
        paragraph.add_run(text)
    style_placeholders(paragraph)


def add_fixed_number_paragraph(document: Document, text: str, keep_with_next: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.paragraph_format.keep_with_next = keep_with_next
    match = re.match(r"^(\d+\.)\s+(.*)$", text)
    if match:
        marker = paragraph.add_run(match.group(1) + " ")
        marker.font.bold = True
        paragraph.add_run(match.group(2))
    else:
        paragraph.add_run(text)
    style_placeholders(paragraph)


def add_checkbox_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    box = paragraph.add_run("☐")
    box.font.name = "Segoe UI Symbol"
    box.font.size = Pt(11)
    paragraph.add_run(text[1:])
    style_placeholders(paragraph)


def style_placeholders(paragraph) -> None:
    full_text = paragraph.text
    if "TÄITA ENNE" in full_text or "[TÄITA]" in full_text:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(RED)
            run.font.bold = True


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    return [rows[0]] + rows[2:]


def table_widths(column_count: int, header: list[str]) -> list[int]:
    if column_count == 2:
        return [3000, 6360]
    if column_count == 3:
        return [2500, 3430, 3430]
    if column_count == 4:
        return [1870, 2810, 2340, 2340]
    if column_count == 5:
        return [2440, 1120, 1320, 2520, 1960]
    each = TABLE_WIDTH_DXA // column_count
    widths = [each] * column_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = table_widths(column_count, rows[0])
    set_table_widths(table, widths)

    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_GREY)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(8.5 if column_count >= 4 else 9)
            if row_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            if "TÄITA ENNE" in value or "[TÄITA]" in value:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(RED)
    repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_body_paragraph(document: Document, text: str, keep_with_next: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.add_run(text)
    style_placeholders(paragraph)


def next_nonblank_line(lines: list[str], start: int) -> str:
    for idx in range(start, len(lines)):
        candidate = lines[idx].strip()
        if candidate:
            return candidate
    return ""


def build_document(source_path: Path, output_path: Path) -> None:
    source_lines = source_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)
    document.core_properties.title = "SotsiaalAI organisatsioonikasutuse raamleping"
    document.core_properties.subject = "Raamleping ja isikuandmete töötlemise kokkulepe"
    document.core_properties.author = "SotsiaalAI OÜ"
    document.core_properties.keywords = "SotsiaalAI, raamleping, piloot, tootmiskasutus, GDPR, andmetöötlus"

    in_summary = False
    first_heading = True
    page_break_before_next = False
    index = 0
    while index < len(source_lines):
        raw = source_lines[index]
        line = raw.strip()

        if not line:
            index += 1
            continue

        if line == "\\newpage":
            page_break_before_next = True
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(source_lines) and source_lines[index].strip().startswith("|"):
                table_lines.append(source_lines[index].strip())
                index += 1
            add_table(document, parse_table(table_lines))
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            if first_heading:
                paragraph = document.add_paragraph(style="Title")
                paragraph.add_run(text)
                set_bottom_border(paragraph)
                first_heading = False
            else:
                paragraph = document.add_paragraph(text, style=f"Heading {min(level, 3)}")
                if page_break_before_next:
                    paragraph.paragraph_format.page_break_before = True
                    page_break_before_next = False
            in_summary = text == "Lepingu lühikokkuvõte"
            index += 1
            continue

        if line == "MUSTAND JURISTI ÜLEVAATUSEKS":
            add_status_paragraph(document, line)
            index += 1
            continue

        if line.startswith("Versioon:") or line.startswith("Dokumendi staatus:") or line.startswith("Asendab pärast"):
            add_metadata_paragraph(document, line.rstrip("  "))
            index += 1
            continue

        if in_summary and re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            paragraph = document.add_paragraph(text, style="List Number")
            if "kõnesalvestus" in text.lower() or "privaat" in text.lower():
                paragraph.runs[0].font.color.rgb = RGBColor.from_string(DARK_BLUE)
            index += 1
            continue

        if re.match(r"^\d+(?:\.\d+)+\.\s", line):
            add_clause_paragraph(document, line)
            index += 1
            continue

        if re.match(r"^[a-zõäöü]\)\s", line, re.IGNORECASE):
            next_line = next_nonblank_line(source_lines, index + 1)
            add_alpha_paragraph(
                document,
                line,
                keep_with_next=bool(re.match(r"^[a-zõäöü]\)\s", next_line, re.IGNORECASE)),
            )
            index += 1
            continue

        if line.startswith("☐"):
            add_checkbox_paragraph(document, line)
            index += 1
            continue

        if re.match(r"^\d+\.\s", line) and not re.match(r"^\d+\.\d+", line):
            next_line = next_nonblank_line(source_lines, index + 1)
            add_fixed_number_paragraph(
                document,
                line,
                keep_with_next=bool(re.match(r"^\d+\.\s", next_line)),
            )
            index += 1
            continue

        next_line = next_nonblank_line(source_lines, index + 1)
        add_body_paragraph(
            document,
            line.rstrip("  "),
            keep_with_next=bool(
                re.match(r"^[a-zõäöü]\)\s", next_line, re.IGNORECASE)
                or re.match(r"^\d+\.\s", next_line)
            ),
        )
        index += 1

    set_repeat_header_footer(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: build_sotsiaalai_framework_docx.py SOURCE.md OUTPUT.docx", file=sys.stderr)
        return 2
    source_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    build_document(source_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
